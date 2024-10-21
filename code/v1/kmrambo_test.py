#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Создано 22 октября 2019 года

Автор: karthick
"""
### Импортируем все необходимые пакеты
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx import *
from docx.text.paragraph import Paragraph
from docx.text.paragraph import Run
import xml.etree.ElementTree as ET
from docx.document import Document as doctwo
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docxcompose.composer import Composer
from docx import Document as Document_compose
import pandas as pd
from xml.etree import ElementTree
from io import StringIO
import io
import csv
import base64
import os  # Импортируем модуль os для работы с путями файлов

# Загружаем файл docx в объект документа. Вы можете указать свой собственный файл docx, изменив путь ниже:
docx_path = r'C:\Users\demchenko\Desktop\SRFauto\test\krambo_test\1.docx'
document = Document(docx_path)

# Получаем директорию, где находится файл docx
output_directory = os.path.dirname(docx_path)

## Эта функция извлекает таблицы и параграфы из объекта документа
def iter_block_items(parent):
    """
    Возвращает каждый параграф и таблицу в *parent* в порядке документа.
    Каждое возвращаемое значение является экземпляром либо Table, либо Paragraph.
    *parent* обычно ссылается на основной объект документа, но также работает для объекта _Cell,
    который сам может содержать параграфы и таблицы.
    """
    if isinstance(parent, doctwo):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Что-то не так")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

# Эта функция извлекает таблицы из объекта документа в формате DataFrame
def read_docx_tables(tab_id=None, **kwargs):
    """
    Парсит таблицы из документа Word (.docx) в Pandas DataFrame(s)

    Параметры:
        tab_id:     парсить одну таблицу с индексом: [tab_id] (начиная с 0).
                    Если [None] - вернуть список DataFrame (парсить все таблицы)

    Возврат: один DataFrame, если tab_id != None или список DataFrame в противном случае
    """
    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in document.tables]
    else:
        try:
            return read_docx_tab(document.tables[tab_id], **kwargs)
        except IndexError:
            print('Ошибка: указанный [tab_id]: {} не существует.'.format(tab_id))
            raise

# Основной код извлечения данных
combined_df = pd.DataFrame(columns=['para_text', 'table_id', 'style'])
table_mod = pd.DataFrame(columns=['string_value', 'table_id'])
image_df = pd.DataFrame(columns=['image_index', 'image_rID', 'image_filename', 'image_base64_string'])
table_list = []
xml_list = []

i = 0
imagecounter = 0

blockxmlstring = ''
for block in iter_block_items(document):
    if 'text' in str(block):
        isappend = False
        
        runboldtext = ''
        for run in block.runs:                        
            if run.bold:
                runboldtext = runboldtext + run.text
                
        style = str(block.style.name)
   
        appendtxt = str(block.text)
        appendtxt = appendtxt.replace("\n", "")
        appendtxt = appendtxt.replace("\r", "")
        tabid = 'Novalue'
        paragraph_split = appendtxt.lower().split()                
        
        isappend = True
        for run in block.runs:
            xmlstr = str(run.element.xml)
            my_namespaces = dict([node for _, node in ElementTree.iterparse(StringIO(xmlstr), events=['start-ns'])])
            root = ET.fromstring(xmlstr) 
            # Проверяем, есть ли изображение в xml элемента. Если да, то извлекаем данные изображения
            if 'pic:pic' in xmlstr:
                xml_list.append(xmlstr)
                for pic in root.findall('.//pic:pic', my_namespaces):
                    cNvPr_elem = pic.find("pic:nvPicPr/pic:cNvPr", my_namespaces)
                    name_attr = cNvPr_elem.get("name")
                    blip_elem = pic.find("pic:blipFill/a:blip", my_namespaces)
                    embed_attr = blip_elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    isappend = True
                    appendtxt = str('Document_Imagefile/' + name_attr + '/' + embed_attr + '/' + str(imagecounter))
                    document_part = document.part
                    image_part = document_part.related_parts[embed_attr]
                    image_base64 = base64.b64encode(image_part._blob)
                    image_base64 = image_base64.decode()                            
                    dftemp = pd.DataFrame({'image_index': [imagecounter], 'image_rID': [embed_attr], 'image_filename': [name_attr], 'image_base64_string': [image_base64]})
                    image_df = pd.concat([image_df, dftemp], ignore_index=True)
                    style = 'Novalue'
                imagecounter = imagecounter + 1
            
    elif 'table' in str(block):
        isappend = True
        style = 'Novalue'
        appendtxt = str(block)
        tabid = i
        dfs = read_docx_tables(tab_id=i)
        dftemp = pd.DataFrame({'para_text': [appendtxt], 'table_id': [i], 'style': [style]})
        table_mod = pd.concat([table_mod, dftemp], ignore_index=True)
        table_list.append(dfs)
        i = i + 1
        
    if isappend:
        dftemp = pd.DataFrame({'para_text': [appendtxt], 'table_id': [tabid], 'style': [style]})
        combined_df = pd.concat([combined_df, dftemp], ignore_index=True)

combined_df = combined_df.reset_index(drop=True)
image_df = image_df.reset_index(drop=True)

# Записываем полученные данные в новый файл в той же директории
output_filename = os.path.join(output_directory, 'output_data.xlsx')
with pd.ExcelWriter(output_filename) as writer:
    combined_df.to_excel(writer, sheet_name='Combined Data', index=False)
    image_df.to_excel(writer, sheet_name='Image Data', index=False)
    table_mod.to_excel(writer, sheet_name='Table Data', index=False)

print(f'Данные успешно записаны в {output_filename}')
