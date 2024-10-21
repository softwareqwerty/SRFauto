import os
import pypandoc
import win32com.client
from docx import Document
from docx.shared import Inches

# Функция для конвертации .doc в .docx
def convert_doc_to_docx(source_path):
    word = win32com.client.Dispatch('Word.Application')
    doc = word.Documents.Open(source_path)
    converted_path = source_path.replace('.doc', '.docx')
    doc.SaveAs(converted_path, FileFormat=16)  # 16 - формат .docx
    doc.Close()
    word.Quit()
    return converted_path

# Функция для копирования элементов (заголовков, параграфов, таблиц и изображений)
def copy_elements(source_doc, target_doc):
    for element in source_doc.element.body:
        if element.tag.endswith('p'):  # Если элемент - параграф
            paragraph = element
            # Копируем текст параграфа
            text = paragraph.text
            if text:
                target_doc.add_paragraph(text)

        elif element.tag.endswith('tbl'):  # Если элемент - таблица
            # Копируем таблицу
            table = element
            new_table = target_doc.add_table(rows=0, cols=len(table.xpath('.//w:tr[1]//w:tc')))
            for row in table.xpath('.//w:tr'):
                new_row = new_table.add_row().cells
                for idx, cell in enumerate(row.xpath('.//w:tc')):
                    text = "".join(cell.itertext()).strip()  # Получаем весь текст в ячейке
                    if text:
                        new_row[idx].text = text

        elif element.tag.endswith('blip'):  # Если элемент - изображение
            # Копируем изображения, если они следуют
            image_data = element.blob
            image_format = element.content_type.split('/')[1]
            temp_image_path = f'temp_image.{image_format}'
            with open(temp_image_path, 'wb') as img_file:
                img_file.write(image_data)
            target_doc.add_picture(temp_image_path, width=Inches(2))  # Задайте нужную ширину
            os.remove(temp_image_path)  # Удаляем временный файл

# Основная функция для переноса данных
def transfer_content(source_path, target_path):
    # Конвертация .doc в .docx
    converted_docx = convert_doc_to_docx(source_path)

    source_doc = Document(converted_docx)
    target_doc = Document()

    copy_elements(source_doc, target_doc)  # Копируем заголовки, текст, таблицы и изображения

    target_doc.save(target_path)

# Указываем путь к исходному и целевому файлам
source_file_path = r"C:\Users\demchenko\Desktop\SRFauto\test\old_docs\1.doc"  # Ваш файл DOC
target_file_path = r"C:\Users\demchenko\Desktop\SRFauto\test\new_docs\output.docx"  # Имя нового DOCX файла

# Выполняем перенос
try:
    transfer_content(source_file_path, target_file_path)
    print("Перенос завершен успешно.")
except Exception as e:
    print(f"Произошла ошибка: {e}")
