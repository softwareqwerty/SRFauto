from docx import Document
import os
from datetime import datetime

def extract_data_from_old_docx(input_path):
    """
    Пример извлечения данных из старого документа .docx.
    Пока возвращает фиксированный текст для замены.
    """
    doc = Document(input_path)
    
    # Пока для примера возвращаем фиксированный текст
    extracted_data = {
        "full_product_name": "Фильтр-регулятор модель XYZ",
        "srf_number": "SRF123456",
        "product_name": "Продукт A",
        "product_measure": "шт.",
        "basic_information": "Это основные сведения об изделии",
        "technical_spec": "Здесь указываются технические характеристики",
        "storage": "Хранить в сухом месте",
        "disposal": "Утилизировать согласно инструкции",
        "packaging_labeling": "Упаковано с осторожностью",
        "exploitation": "Эксплуатация в температурном диапазоне от -20 до +40°C",
        "rev_number": "Ревизия 1"
    }
    
    return extracted_data

def copy_run_format(from_run, to_run):
    """
    Копирует форматирование с одного run на другой.
    """
    to_run.font.size = from_run.font.size
    to_run.font.bold = from_run.font.bold
    to_run.font.italic = from_run.font.italic
    to_run.font.underline = from_run.font.underline
    to_run.font.color.rgb = from_run.font.color.rgb
    to_run.font.name = from_run.font.name

def replace_markers_in_paragraph(paragraph, data):
    """
    Заменяет маркеры в абзаце и применяет форматирование к вставляемому тексту.
    """
    for key, value in data.items():
        placeholder = f"{{{{{key}}}}}"  # Формируем шаблон маркера, например, {{date}}

        # Обходим все runs в параграфе
        for run in paragraph.runs:
            if placeholder in run.text:
                # Сохраняем форматирование маркера
                original_format = run.font

                # Разделяем текст на части: до маркера и после маркера
                parts = run.text.split(placeholder)

                # Заменяем текст до маркера
                run.text = parts[0]  # Устанавливаем текст до маркера

                # Добавляем новый run для вставляемого текста с сохраненным форматированием
                new_run = paragraph.add_run(value)
                copy_run_format(original_format, new_run)  # Копируем форматирование

                # Если есть текст после маркера, добавляем его
                if len(parts) > 1 and parts[1]:
                    run = paragraph.add_run(parts[1])
                    copy_run_format(original_format, run)  # Копируем форматирование

                break  # Прекращаем обработку после первой замены

def apply_data_to_template(template_path, output_path, data):
    """
    Заменяет маркеры в шаблоне извлеченными данными и сохраняет новый документ.
    """
    doc = Document(template_path)
    
    # Проходим по всем абзацам в документе
    for para in doc.paragraphs:
        replace_markers_in_paragraph(para, data)
    
    # Проходим по всем таблицам в документе
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_markers_in_paragraph(para, data)
    
    # Обработка верхнего и нижнего колонтитулов
    for section in doc.sections:
        header = section.header
        footer = section.footer
        
        # Проходим по абзацам верхнего колонтитула
        for para in header.paragraphs:
            replace_markers_in_paragraph(para, data)
        
        # Проходим по таблицам в верхнем колонтитуле
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_markers_in_paragraph(para, data)
        
        # Проходим по абзацам нижнего колонтитула
        for para in footer.paragraphs:
            replace_markers_in_paragraph(para, data)
        
        # Проходим по таблицам в нижнем колонтитуле
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_markers_in_paragraph(para, data)
    
    # Сохраняем измененный документ
    doc.save(output_path)

def process_files_with_template(input_folder, template_path, output_folder):
    """
    Основная функция обработки:
    1. Чтение данных из старого .docx.
    2. Вставка данных в шаблон.
    3. Сохранение нового документа.
    """
    # Получение текущей даты
    current_date = datetime.now().strftime("%d.%m.%Y")
    
    # Создаем выходную папку, если её нет
    os.makedirs(output_folder, exist_ok=True)
    
    for filename in os.listdir(input_folder):
        if filename.endswith(".docx"):
            input_path = os.path.join(input_folder, filename)
            output_path = os.path.join(output_folder, "new_" + filename)
            
            # Извлечение данных из старого документа
            extracted_data = extract_data_from_old_docx(input_path)
            
            # Добавляем текущую дату в данные
            extracted_data["date"] = current_date
            
            # Применение данных к шаблону и сохранение результата
            print(f"Обработка файла {filename} по шаблону...")
            apply_data_to_template(template_path, output_path, extracted_data)
            
            print(f"Обработан и сохранен файл: {output_path}")

# Указываем пути к папкам и файлам
input_folder = r"C:\Users\demchenko\Desktop\SRFauto\test\old_docs"
template_path = r"C:\Users\demchenko\Desktop\SRFauto\test\template\3.docx"
output_folder = r"C:\Users\demchenko\Desktop\SRFauto\test\new_docs"

# Запуск обработки
process_files_with_template(input_folder, template_path, output_folder)
