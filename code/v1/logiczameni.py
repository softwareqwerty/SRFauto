from docx import Document
import os
import win32com.client as win32
from datetime import datetime

def convert_doc_to_docx(input_path, output_path):
    """
    Конвертирует .doc файл в .docx.
    """
    word = win32.Dispatch("Word.Application")
    try:
        doc = word.Documents.Open(input_path)
        doc.SaveAs(output_path, FileFormat=16)  # 16 - формат .docx
        doc.Close()
    except Exception as e:
        print(f"Ошибка при конвертации {input_path}: {e}")
    finally:
        word.Quit()


import re  # Импортируем модуль для работы с регулярными выражениями


def extract_data_from_old_docx(input_path):
    """
    Извлекает текстовые данные из старого документа .docx.
    Возвращает пример данных для замены маркеров.
    """
    doc = Document(input_path)

    # Пример извлечения данных
    extracted_data = {
        "full_product_name": "",  # Изначально пусто
        "srf_number": "",  # Изначально пусто
        "product_name": "Продукт A",
        "product_measure": "шт.",
        "basic_information": "",  # Изначально пусто
        "technical_spec": "",  # Изначально пусто
        "storage": "",  # Изначально пусто
        "disposal": "",  # Изначально пусто
        "packaging_labeling": "",  # Изначально пусто
        "exploitation": "",  # Изначально пусто
        "rev_number": "",  # Изначально пусто
        "date": datetime.now().strftime("%d.%m.%Y")  # Добавляем текущую дату
    }

    current_section = None  # Переменная для отслеживания текущей секции
    text_buffer = []  # Буфер для хранения текста секции

    def save_section():
        """Сохраняет собранный текст для текущей секции в extracted_data."""
        if current_section and text_buffer:
            extracted_data[current_section] = "\n".join(text_buffer).strip()
            print(f"Сохранён текст для '{current_section}': '{extracted_data[current_section]}'")  # Отладочное сообщение
            text_buffer.clear()

    # Логика для извлечения srf_number
    srf_pattern = re.compile(r"\bSRF[-\w]+\b", re.IGNORECASE)  # Регулярное выражение для поиска "SRF" + буквы и цифры
    found_srf = False  # Флаг для отслеживания, нашли ли мы уже srf_number

    # Логика для извлечения full_product_name
    passport_pattern = re.compile(r"\bпаспорт\b", re.IGNORECASE)  # Регулярное выражение для поиска "паспорт"
    found_passport = False  # Флаг для отслеживания, нашли ли мы уже паспорт

    # Логика для извлечения rev_number
    rev_pattern = re.compile(r"\bрев\w*", re.IGNORECASE)  # Регулярное выражение для поиска "рев" + буквы
    found_rev = False  # Флаг для отслеживания, нашли ли мы уже rev_number

    # Проверка таблиц в верхнем колонтитуле на наличие srf_number и паспорта
    for section in doc.sections:
        header = section.header
        # Проверяем каждую таблицу в верхнем колонтитуле
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    print(f"Проверяем ячейку колонтитула: '{cell.text}'")  # Отладочное сообщение
                    
                    # Поиск номера SRF
                    if not found_srf:
                        match = srf_pattern.search(cell.text)
                        if match:
                            extracted_data["srf_number"] = match.group(0)
                            found_srf = True
                            print(f"Найден SRF номер в таблице колонтитула: '{extracted_data['srf_number']}'")  # Отладочное сообщение

                    # Поиск слова "паспорт"
                    if not found_passport:
                        match = passport_pattern.search(cell.text)
                        if match:
                            # Удаляем слово "паспорт" из текста
                            full_product_name = cell.text.replace(match.group(0), "").strip()
                            extracted_data["full_product_name"] = full_product_name
                            found_passport = True
                            print(f"Найден паспорт в таблице колонтитула: '{extracted_data['full_product_name']}'")  # Отладочное сообщение
                    
                    # Поиск слова "рев" для rev_number
                    if not found_rev:
                        match = rev_pattern.search(cell.text)
                        if match:
                            # Проверяем, есть ли символ после "рев"
                            after_rev = cell.text[match.end():].strip()  # Текст после найденного "рев"
                            if after_rev and after_rev[0].isalpha():  # Если есть буква
                                extracted_data["rev_number"] = after_rev[0]  # Берем первую букву
                                found_rev = True
                                print(f"Найден rev_number из ячейки: '{extracted_data['rev_number']}'")  # Отладочное сообщение
                            else:
                                # Если буквы нет, нужно проверять следующую ячейку
                                print("Нет буквы после 'рев', будем искать в следующей ячейке.")  # Отладочное сообщение
                                

                                

    # Проходим по всем абзацам в основном документе
    for para in doc.paragraphs:
        print(f"Обрабатываем абзац: '{para.text}', стиль: '{para.style.name}'")  # Отладочное сообщение

        # Проверка на наличие SRF номера, если он еще не найден
        if not found_srf:
            match = srf_pattern.search(para.text)
            if match:
                extracted_data["srf_number"] = match.group(0)
                found_srf = True
                print(f"Найден SRF номер: '{extracted_data['srf_number']}'")  # Отладочное сообщение

        # Проверка заголовков для других секций
        if para.style.name.startswith('Heading'):
            # Сначала сохраняем текст предыдущей секции перед переключением
            save_section()

            # Проверка, содержит ли заголовок нужные ключевые слова
            if "хранени" in para.text.lower():
                current_section = "storage"
                print("Обнаружен заголовок 'Хранени'. Начинаем сбор текста.")  # Отладочное сообщение
            elif "изделии" in para.text.lower():
                current_section = "basic_information"
                print("Обнаружен заголовок 'Изделии'. Начинаем сбор текста.")  # Отладочное сообщение
            elif "характеристики" in para.text.lower():
                current_section = "technical_spec"
                print("Обнаружен заголовок 'Характеристики'. Начинаем сбор текста.")  # Отладочное сообщение
            elif "утилизаци" in para.text.lower():
                current_section = "disposal"
                print("Обнаружен заголовок 'Утилизаци'. Начинаем сбор текста.")  # Отладочное сообщение
            elif "упаковк" in para.text.lower():  # Проверка на частичное совпадение
                current_section = "packaging_labeling"
                print("Обнаружен заголовок 'Упаковка'. Начинаем сбор текста.")  # Отладочное сообщение
            elif "монтаж" in para.text.lower():
                current_section = "exploitation"
                print("Обнаружен заголовок 'Монтаж'. Начинаем сбор текста.")  # Отладочное сообщение
            else:
                current_section = None  # Сбрасываем текущую секцию, если заголовок не соответствует

            continue  # Переходим к следующему абзацу

        # Если мы находимся в секции, собираем текст
        if current_section:
            if para.text.strip() == "":  # Если пустой абзац, пропускаем его
                print("Пустой абзац, пропускаем.")  # Отладочное сообщение
                continue
            
            # Добавляем текст абзаца в буфер
            print(f"Добавляем текст в '{current_section}': '{para.text}'")  # Отладочное сообщение
            text_buffer.append(para.text)

    # Сохраняем текст для последней секции, если она была активной
    save_section()

    # Выводим итоговые значения для всех полей
    for key, value in extracted_data.items():
        print(f"ИТОГО {key}: '{value}'")  # Итоговые значения для всех полей

    return extracted_data




def replace_markers_in_paragraph(paragraph, data):
    """ 
    Заменяет маркеры в абзаце, сохраняя форматирование.
    """
    for key, value in data.items():
        placeholder = f"{{{{{key}}}}}"  # Формируем шаблон маркера
        full_text = ''.join(run.text for run in paragraph.runs)  # Получаем весь текст абзаца

        if placeholder in full_text:
            print(f"Найден маркер: '{placeholder}' в абзаце: '{full_text}'")  # Отладочное сообщение
            
            # Считываем форматирование первой буквы маркера
            format_run = None
            for run in paragraph.runs:
                if placeholder[0] in run.text:
                    format_run = run
                    break

            if format_run:
                print(f"Используем форматирование из 'Run': '{format_run.text}'")  # Отладочное сообщение
                
                # Получаем атрибуты форматирования
                size = format_run.font.size
                bold = format_run.font.bold
                italic = format_run.font.italic
                color = format_run.font.color.rgb
                font_name = format_run.font.name

                print(f"Форматирование из первой буквы маркера: размер={size}, жирный={bold}, курсив={italic}, цвет={color}, шрифт={font_name}")  # Отладочное сообщение
            else:
                print(f"Не удалось найти форматирование для маркера '{placeholder}', используем дефолтные значения.")

            # Очищаем все старые `Run` в абзаце
            paragraph.clear()
            
            # Создаем новый `Run` с заменённым текстом и применяем форматирование
            new_run = paragraph.add_run(full_text.replace(placeholder, value))
            if format_run:
                new_run.font.size = size or 12  # Используем 12, если `None`
                new_run.font.bold = bold
                new_run.font.italic = italic
                if color:
                    new_run.font.color.rgb = color
                new_run.font.name = font_name or "Arial"
            
            print(f"Заменено '{placeholder}' на '{new_run.text}' с применением форматирования.")  # Сообщение о замене

def apply_data_to_template(template_path, output_path, data):
    """
    Заменяет маркеры в шаблоне извлеченными данными и сохраняет новый документ.
    """
    doc = Document(template_path)

    # Заменяем маркеры в абзацах
    for para in doc.paragraphs:
        replace_markers_in_paragraph(para, data)

    # Заменяем маркеры в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_markers_in_paragraph(para, data)

    # Обработка колонтитулов
    for section in doc.sections:
        header = section.header
        footer = section.footer
        
        # Обработка верхнего колонтитула
        for para in header.paragraphs:
            replace_markers_in_paragraph(para, data)
        
        # Обработка нижнего колонтитула
        for para in footer.paragraphs:
            replace_markers_in_paragraph(para, data)

        # Обработка таблиц в верхнем колонтитуле
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_markers_in_paragraph(para, data)
        
        # Обработка таблиц в нижнем колонтитуле
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_markers_in_paragraph(para, data)

    # Сохраняем измененный документ
    doc.save(output_path)

def process_files_with_template(input_folder, template_path, output_folder):
    """
    Основная функция обработки:
    1. Конвертация .doc в .docx.
    2. Замена маркеров в шаблоне.
    3. Сохранение готового документа.
    """
    # Создаем выходную папку, если её нет
    os.makedirs(output_folder, exist_ok=True)

    for filename in os.listdir(input_folder):
        if filename.endswith(".doc"):
            input_path = os.path.join(input_folder, filename)
            temp_docx_path = os.path.join(input_folder, filename.replace(".doc", ".docx"))
            output_path = os.path.join(output_folder, "new_" + filename.replace(".doc", ".docx"))

            # Конвертация .doc в .docx
            print(f"Конвертация файла {filename} в .docx...")
            convert_doc_to_docx(input_path, temp_docx_path)

            # Извлечение данных из старого документа
            extracted_data = extract_data_from_old_docx(temp_docx_path)

            # Применение данных к шаблону и сохранение результата
            print(f"Обработка файла {filename} по шаблону...")
            apply_data_to_template(template_path, output_path, extracted_data)

            # Удаление временного .docx файла
            os.remove(temp_docx_path)
            print(f"Обработан и сохранен файл: {output_path}")

# Указываем пути к папкам и файлам
input_folder = r"C:\Users\demchenko\Desktop\SRFauto\test\old_docs"
template_path = r"C:\Users\demchenko\Desktop\SRFauto\test\template\3.docx"
output_folder = r"C:\Users\demchenko\Desktop\SRFauto\test\new_docs"

# Запуск обработки
process_files_with_template(input_folder, template_path, output_folder)
