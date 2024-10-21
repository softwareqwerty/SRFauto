#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Создано 22 октября 2019 года

Автор: karthick
"""
### Импортируем все необходимые пакеты
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx import *
from docx.text.paragraph import Paragraph
from docx.text.paragraph import Run
import xml.etree.ElementTree as ET
from docx.document import Document as doctwo
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.shared import Pt
import pandas as pd
from xml.etree import ElementTree
from io import StringIO
import io
import csv
import base64
import os  # Импортируем модуль os для работы с путями файлов

# Загружаем файл docx в объект документа. Вы можете указать свой собственный файл docx, изменив путь ниже:
docx_path = r'C:\Users\demchenko\Desktop\SRFauto\test\krambo_test\1.docx'
document = Document(docx_path)

# Получаем директорию, где находится файл docx
output_directory = os.path.dirname(docx_path)

## Эта функция извлекает таблицы и параграфы из объекта документа
def iter_block_items(parent):
    if isinstance(parent, doctwo):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Что-то не так")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

# Эта функция извлекает таблицы из объекта документа в формате DataFrame
def read_docx_tables(tab_id=None, **kwargs):
    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in document.tables]
    else:
        try:
            return read_docx_tab(document.tables[tab_id], **kwargs)
        except IndexError:
            print('Ошибка: указанный [tab_id]: {} не существует.'.format(tab_id))
            raise

# ... Предыдущие импорты и код остаются без изменений ...

import os
import docx
from docx import Document
from docx.shared import Inches
import pandas as pd
import xml.etree.ElementTree as ET
from io import StringIO
import base64

# Функция для декодирования Base64 строки и сохранения изображения
def decode_image(base64_string, output_path):
    image_data = base64.b64decode(base64_string)
    with open(output_path, 'wb') as f:
        f.write(image_data)
    return output_path

# Основной код извлечения данных
combined_df = pd.DataFrame(columns=['para_text', 'table_id', 'style'])
table_mod = pd.DataFrame(columns=['string_value', 'table_id'])
image_df = pd.DataFrame(columns=['image_index', 'image_rID', 'image_filename', 'image_base64_string'])
table_list = []
table_image_data = []  # Объявление списка для хранения информации о изображениях в таблицах
xml_list = []

i = 0
imagecounter = 0

blockxmlstring = ''
for block in iter_block_items(document):
    if 'text' in str(block):
        isappend = False
        runboldtext = ''
        for run in block.runs:                        
            if run.bold:
                runboldtext = runboldtext + run.text
                
        style = str(block.style.name)
        appendtxt = str(block.text).replace("\n", "").replace("\r", "")
        tabid = 'Novalue'
        
        isappend = True
        for run in block.runs:
            xmlstr = str(run.element.xml)
            print(f"Обработка XML строки: {xmlstr[:100]}...")  # Отладочное сообщение

            # Проверка на пустую строку перед парсингом
            if not xmlstr.strip():
                print("Предупреждение: пустая строка для xmlstr. Пропускаем.")  # Отладочное сообщение
                continue
            
            try:
                my_namespaces = dict([node for _, node in ET.iterparse(StringIO(xmlstr), events=['start-ns'])])
                root = ET.fromstring(xmlstr)
            except ET.ParseError as e:
                print(f"Ошибка парсинга XML: {e} в строке: {xmlstr[:100]}... Пропускаем.")  # Отладочное сообщение
                continue
            
            # Проверяем, есть ли изображение в xml элемента. Если да, то извлекаем данные изображения
            if 'pic:pic' in xmlstr:
                xml_list.append(xmlstr)
                for pic in root.findall('.//pic:pic', my_namespaces):
                    cNvPr_elem = pic.find("pic:nvPicPr/pic:cNvPr", my_namespaces)
                    name_attr = cNvPr_elem.get("name")
                    blip_elem = pic.find("pic:blipFill/a:blip", my_namespaces)
                    embed_attr = blip_elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    isappend = True
                    appendtxt = str('Document_Imagefile/' + name_attr + '/' + embed_attr + '/' + str(imagecounter))
                    document_part = document.part
                    image_part = document_part.related_parts[embed_attr]
                    image_base64 = base64.b64encode(image_part._blob)
                    image_base64 = image_base64.decode()                            
                    dftemp = pd.DataFrame({'image_index': [imagecounter], 'image_rID': [embed_attr], 'image_filename': [name_attr], 'image_base64_string': [image_base64]})
                    image_df = pd.concat([image_df, dftemp], ignore_index=True)
                    table_image_data.append({
                        'image_filename': name_attr,
                        'image_base64_string': image_base64,
                        'table_id': i,  # Сохраняем ID таблицы
                        'row_index': None,  # Пока не знаем, в какой ячейке, обновим позже
                        'cell_index': None  # Пока не знаем, в какой ячейке, обновим позже
                    })
                    style = 'Novalue'
                imagecounter += 1
            
    elif 'table' in str(block):
        print(f'Обрабатывается таблица {i}')  # Отладочное сообщение
        isappend = True
        style = 'Novalue'
        appendtxt = str(block)
        tabid = i
        dfs = read_docx_tables(tab_id=i)
        
        # Сохранение информации о таблице
        dftemp = pd.DataFrame({'para_text': [appendtxt], 'table_id': [i], 'style': [style]})
        table_mod = pd.concat([table_mod, dftemp], ignore_index=True)
        table_list.append(dfs)
        
        # Извлечение изображений из таблицы
        for r_idx, row in dfs.iterrows():
            for c_idx, cell in enumerate(row):
                # Проверка, есть ли изображения в ячейках
                xmlstr = str(cell)  # Приводим к строке для обработки
                print(f"Обработка XML строки для ячейки ({r_idx}, {c_idx}): {xmlstr[:100]}...")  # Выводим начало xmlstr для диагностики
                
                if not xmlstr.strip():
                    print(f"Предупреждение: ячейка {r_idx}, {c_idx} пустая. Пропускаем.")  # Отладочное сообщение
                    continue

                try:
                    my_namespaces = dict([node for _, node in ET.iterparse(StringIO(xmlstr), events=['start-ns'])])
                    root = ET.fromstring(xmlstr)
                except ET.ParseError as e:
                    print(f"Ошибка парсинга XML в ячейке ({r_idx}, {c_idx}): {e}. Пропускаем.")  # Отладочное сообщение
                    continue

                # Проверяем наличие изображений в ячейках
                if 'pic:pic' in xmlstr:
                    for pic in root.findall('.//pic:pic', my_namespaces):
                        cNvPr_elem = pic.find("pic:nvPicPr/pic:cNvPr", my_namespaces)
                        name_attr = cNvPr_elem.get("name")
                        blip_elem = pic.find("pic:blipFill/a:blip", my_namespaces)
                        embed_attr = blip_elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        
                        # Получаем изображение
                        image_part = document.part.related_parts[embed_attr]
                        image_base64 = base64.b64encode(image_part._blob)
                        image_base64 = image_base64.decode()

                        # Сохранение информации о изображении с указанием позиции
                        table_image_data.append({
                            'image_filename': name_attr,
                            'image_base64_string': image_base64,
                            'table_id': i,
                            'row_index': r_idx,
                            'cell_index': c_idx
                        })
                    print(f'Изображение {name_attr} найдено в таблице {i}, ячейка {c_idx}, строка {r_idx}')  # Отладочное сообщение

        i += 1  # Увеличиваем ID таблицы на 1

    if isappend:
        combined_df = pd.concat([combined_df, pd.DataFrame({'para_text': [appendtxt], 'table_id': [tabid], 'style': [style]})], ignore_index=True)

# Вставка изображений обратно в таблицы и создание нового документа
new_doc = Document()
for table_index, table_data in enumerate(table_list):
    new_table = new_doc.add_table(rows=len(table_data), cols=len(table_data.columns))
    new_table.style = document.tables[table_index].style  # Устанавливаем стиль таблицы
    
    for r_idx, row in table_data.iterrows():
        for c_idx, cell in enumerate(row):
            new_cell = new_table.cell(r_idx, c_idx)
            new_cell.text = str(cell)
            
            # Проверка, есть ли изображение для текущей ячейки
            for img_data in table_image_data:
                if img_data['table_id'] == table_index and img_data['row_index'] == r_idx and img_data['cell_index'] == c_idx:
                    image_filename = img_data['image_filename']
                    image_base64_string = img_data['image_base64_string']
                    output_image_path = f"temp_{image_filename}"
                    decode_image(image_base64_string, output_image_path)
                    new_cell.paragraphs[0].add_run().add_picture(output_image_path, width=Inches(1.0))
                    os.remove(output_image_path)  # Удаляем временный файл

new_doc.save('new_document_with_images.docx')
