from docx import Document
import os
import win32com.client as win32
from datetime import datetime

def convert_doc_to_docx(input_path, output_path):
    """
    Конвертирует .doc файл в .docx.
    """
    word = win32.Dispatch("Word.Application")
    try:
        doc = word.Documents.Open(input_path)
        doc.SaveAs(output_path, FileFormat=16)  # 16 - формат .docx
        doc.Close()
    except Exception as e:
        print(f"Ошибка при конвертации {input_path}: {e}")
    finally:
        word.Quit()


import re  # Импортируем модуль для работы с регулярными выражениями


from docx import Document
import re
from datetime import datetime

def extract_data_from_old_docx(input_path):
    """
    Извлекает текстовые данные из старого документа .docx.
    Возвращает пример данных для замены маркеров.
    """
    doc = Document(input_path)

    # Инициализируем структуру для извлеченных данных
    extracted_data = {
        "full_product_name": "",
        "srf_number": "",
        "product_name": "",
        "product_measure": "",
        "basic_information": "",
        "technical_spec": "",
        "storage": "",
        "disposal": "",
        "packaging_labeling": "",
        "exploitation": "",
        "rev_number": "",
        "date": datetime.now().strftime("%d.%m.%Y")
    }

    # Регулярные выражения для поиска нужных данных
    srf_pattern = re.compile(r"\bSRF[-\w]+\b", re.IGNORECASE)
    passport_pattern = re.compile(r"\bпаспорт\b", re.IGNORECASE)
    rev_pattern = re.compile(r"\bрев\w*[:.,;\s]*", re.IGNORECASE)

    current_section = None
    text_buffer = []

    def save_section():
        """Сохраняет собранный текст для текущей секции."""
        if current_section and text_buffer:
            extracted_data[current_section] = "\n".join(text_buffer).strip()
            text_buffer.clear()

    # Поиск информации в колонтитулах
    for section in doc.sections:
        header = section.header
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Поиск номера SRF
                    if not extracted_data["srf_number"]:
                        match = srf_pattern.search(cell.text)
                        if match:
                            extracted_data["srf_number"] = match.group(0)

                    # Поиск полного наименования продукта
                    if not extracted_data["full_product_name"]:
                        match = passport_pattern.search(cell.text)
                        if match:
                            extracted_data["full_product_name"] = cell.text.replace(match.group(0), "").strip()

                    # Поиск номера ревизии
                    if not extracted_data["rev_number"]:
                        match = rev_pattern.search(cell.text)
                        if match:
                            after_rev = cell.text[match.end():].strip()
                            if after_rev:
                                extracted_data["rev_number"] = after_rev[0]  # Используем первый символ после "рев"

    # Обработка текста в документе для извлечения информации по секциям
    for para in doc.paragraphs:
        # Если это заголовок секции, определяем текущую секцию
        if para.style.name.startswith('Heading'):
            save_section()

            if "хранени" in para.text.lower():
                current_section = "storage"
            elif "изделии" in para.text.lower():
                current_section = "basic_information"
            elif "характеристики" in para.text.lower():
                current_section = "technical_spec"
            elif "утилизаци" in para.text.lower():
                current_section = "disposal"
            elif "упаковк" in para.text.lower():
                current_section = "packaging_labeling"
            elif "монтаж" in para.text.lower():
                current_section = "exploitation"
            else:
                current_section = None
            continue

        # Сбор текста для текущей секции
        if current_section:
            text_buffer.append(para.text.strip())

        # Поиск SRF номера в тексте документа, если он еще не найден
        if not extracted_data["srf_number"]:
            match = srf_pattern.search(para.text)
            if match:
                extracted_data["srf_number"] = match.group(0)

    # Сохраняем текст для последней секции
    save_section()

    # Поиск информации о продукте и его измерениях в таблицах
    for table in doc.tables:
        for row in table.rows:
            first_cell_text = row.cells[0].text.strip()

            if "," in first_cell_text:
                parts = first_cell_text.split(",", 1)
                extracted_data["product_name"] = parts[0].strip()
                extracted_data["product_measure"] = parts[1].strip()
                break

    return extracted_data






def replace_markers_in_paragraph(paragraph, data):
    """ 
    Заменяет маркеры в абзаце, сохраняя форматирование.
    """
    for key, value in data.items():
        placeholder = f"{{{{{key}}}}}"  # Формируем шаблон маркера
        full_text = ''.join(run.text for run in paragraph.runs)  # Получаем весь текст абзаца

        if placeholder in full_text:
            print(f"Найден маркер: '{placeholder}' в абзаце: '{full_text}'")  # Отладочное сообщение
            
            # Считываем форматирование первой буквы маркера
            format_run = None
            for run in paragraph.runs:
                if placeholder[0] in run.text:
                    format_run = run
                    break

            if format_run:
                print(f"Используем форматирование из 'Run': '{format_run.text}'")  # Отладочное сообщение
                
                # Получаем атрибуты форматирования
                size = format_run.font.size
                bold = format_run.font.bold
                italic = format_run.font.italic
                color = format_run.font.color.rgb
                font_name = format_run.font.name

                print(f"Форматирование из первой буквы маркера: размер={size}, жирный={bold}, курсив={italic}, цвет={color}, шрифт={font_name}")  # Отладочное сообщение
            else:
                print(f"Не удалось найти форматирование для маркера '{placeholder}', используем дефолтные значения.")

            # Очищаем все старые `Run` в абзаце
            paragraph.clear()
            
            # Создаем новый `Run` с заменённым текстом и применяем форматирование
            new_run = paragraph.add_run(full_text.replace(placeholder, value))
            if format_run:
                new_run.font.size = size or 12  # Используем 12, если `None`
                new_run.font.bold = bold
                new_run.font.italic = italic
                if color:
                    new_run.font.color.rgb = color
                new_run.font.name = font_name or "Arial"
            
            print(f"Заменено '{placeholder}' на '{new_run.text}' с применением форматирования.")  # Сообщение о замене

def apply_data_to_template(template_path, output_path, data):
    """
    Заменяет маркеры в шаблоне извлеченными данными и сохраняет новый документ.
    """
    doc = Document(template_path)

    # Заменяем маркеры в абзацах
    for para in doc.paragraphs:
        replace_markers_in_paragraph(para, data)

    # Заменяем маркеры в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_markers_in_paragraph(para, data)

    # Обработка колонтитулов
    for section in doc.sections:
        header = section.header
        footer = section.footer
        
        # Обработка верхнего колонтитула
        for para in header.paragraphs:
            replace_markers_in_paragraph(para, data)
        
        # Обработка нижнего колонтитула
        for para in footer.paragraphs:
            replace_markers_in_paragraph(para, data)

        # Обработка таблиц в верхнем колонтитуле
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_markers_in_paragraph(para, data)
        
        # Обработка таблиц в нижнем колонтитуле
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_markers_in_paragraph(para, data)

    # Сохраняем измененный документ
    doc.save(output_path)

def process_files_with_template(input_folder, template_path, output_folder):
    """
    Основная функция обработки:
    1. Конвертация .doc в .docx.
    2. Замена маркеров в шаблоне.
    3. Сохранение готового документа.
    """
    # Создаем выходную папку, если её нет
    os.makedirs(output_folder, exist_ok=True)

    for filename in os.listdir(input_folder):
        if filename.endswith(".doc"):
            input_path = os.path.join(input_folder, filename)
            temp_docx_path = os.path.join(input_folder, filename.replace(".doc", ".docx"))
            output_path = os.path.join(output_folder, "new_" + filename.replace(".doc", ".docx"))

            # Конвертация .doc в .docx
            print(f"Конвертация файла {filename} в .docx...")
            convert_doc_to_docx(input_path, temp_docx_path)

            # Извлечение данных из старого документа
            extracted_data = extract_data_from_old_docx(temp_docx_path)

            # Применение данных к шаблону и сохранение результата
            print(f"Обработка файла {filename} по шаблону...")
            apply_data_to_template(template_path, output_path, extracted_data)

            # Удаление временного .docx файла
            os.remove(temp_docx_path)
            print(f"Обработан и сохранен файл: {output_path}")


# Указываем пути к папкам и файлам
input_folder = r"C:\Users\demchenko\Desktop\SRFauto\test\old_docs"
template_path = r"C:\Users\demchenko\Desktop\SRFauto\test\template\3.docx"
output_folder = r"C:\Users\demchenko\Desktop\SRFauto\test\new_docs"

# Запуск обработки
process_files_with_template(input_folder, template_path, output_folder)
