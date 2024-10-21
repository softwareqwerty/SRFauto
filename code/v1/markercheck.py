import os
import win32com.client as win32
from datetime import datetime
from docx import Document

def convert_doc_to_docx(input_path, output_path):
    """
    Конвертирует .doc файл в .docx.
    """
    word = win32.Dispatch("Word.Application")
    try:
        doc = word.Documents.Open(input_path)
        doc.SaveAs(output_path, FileFormat=16)  # 16 - формат .docx
        doc.Close()
    except Exception as e:
        print(f"Ошибка при конвертации {input_path}: {e}")
    finally:
        word.Quit()

def extract_data_from_old_docx(input_path):
    """
    Извлекает текстовые данные из старого документа .docx.
    Возвращает пример данных для замены маркеров.
    """
    doc = Document(input_path)
    
    # Пример извлечения данных
    extracted_data = {
        "full_product_name": "Фильтр-регулятор модель XYZ",
        "srf_number": "SRF123456",
        "product_name": "Продукт A",
        "product_measure": "шт.",
        "basic_information": "Это основные сведения об изделии",
        "technical_spec": "Здесь указываются технические характеристики",
        "storage": "Хранить в сухом месте",
        "disposal": "Утилизировать согласно инструкции",
        "packaging_labeling": "Упаковано с осторожностью",
        "exploitation": "Эксплуатация в температурном диапазоне от -20 до +40°C",
        "rev_number": "Ревизия 1"
    }
    
    return extracted_data

def replace_markers_in_paragraph(paragraph, data):
    """ 
    Заменяет маркеры в абзаце. 
    """
    for key, value in data.items():
        placeholder = f"{{{{{key}}}}}"  # Формируем шаблон маркера
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, value)

def apply_data_to_template(template_path, output_path, data):
    """
    Заменяет маркеры в шаблоне извлеченными данными и сохраняет новый документ.
    """
    doc = Document(template_path)

    # Заменяем маркеры в абзацах
    for para in doc.paragraphs:
        replace_markers_in_paragraph(para, data)

    # Заменяем маркеры в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_markers_in_paragraph(para, data)

    # Обработка колонтитулов
    for section in doc.sections:
        header = section.header
        footer = section.footer
        
        # Обработка верхнего колонтитула
        for para in header.paragraphs:
            replace_markers_in_paragraph(para, data)
        
        # Обработка нижнего колонтитула
        for para in footer.paragraphs:
            replace_markers_in_paragraph(para, data)

        # Обработка таблиц в верхнем колонтитуле
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_markers_in_paragraph(para, data)
        
        # Обработка таблиц в нижнем колонтитуле
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_markers_in_paragraph(para, data)

    # Сохраняем измененный документ
    doc.save(output_path)

def process_files_with_template(input_folder, template_path, output_folder):
    """
    Основная функция обработки:
    1. Конвертация .doc в .docx.
    2. Замена маркеров в шаблоне.
    3. Сохранение готового документа.
    """
    # Получение текущей даты
    current_date = datetime.now().strftime("%d.%m.%Y")

    # Создаем выходную папку, если её нет
    os.makedirs(output_folder, exist_ok=True)

    for filename in os.listdir(input_folder):
        if filename.endswith(".doc"):
            input_path = os.path.join(input_folder, filename)
            temp_docx_path = os.path.join(input_folder, filename.replace(".doc", ".docx"))
            output_path = os.path.join(output_folder, "new_" + filename.replace(".doc", ".docx"))

            # Конвертация .doc в .docx
            print(f"Конвертация файла {filename} в .docx...")
            convert_doc_to_docx(input_path, temp_docx_path)

            # Извлечение данных из старого документа
            extracted_data = extract_data_from_old_docx(temp_docx_path)

            # Добавляем текущую дату в данные
            extracted_data["date"] = current_date

            # Применение данных к шаблону и сохранение результата
            print(f"Обработка файла {filename} по шаблону...")
            apply_data_to_template(template_path, output_path, extracted_data)

            # Удаление временного .docx файла
            os.remove(temp_docx_path)
            print(f"Обработан и сохранен файл: {output_path}")

# Указываем пути к папкам и файлам
input_folder = r"C:\Users\demchenko\Desktop\SRFauto\test\old_docs"
template_path = r"C:\Users\demchenko\Desktop\SRFauto\test\template\3.docx"
output_folder = r"C:\Users\demchenko\Desktop\SRFauto\test\new_docs"

# Запуск обработки
process_files_with_template(input_folder, template_path, output_folder)
