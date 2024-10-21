import os
import pypandoc
import win32com.client
from docx import Document
from docx.shared import Inches

# Функция для конвертации .doc в .docx
def convert_doc_to_docx(source_path):
    word = win32com.client.Dispatch('Word.Application')
    doc = word.Documents.Open(source_path)
    converted_path = source_path.replace('.doc', '.docx')
    doc.SaveAs(converted_path, FileFormat=16)  # 16 - формат .docx
    doc.Close()
    word.Quit()
    return converted_path

# Функция для копирования изображений
def copy_images(source_doc, target_doc):
    for rel in source_doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_format = rel.target_part.content_type.split('/')[1]  # Получаем формат изображения
            # Сохраняем изображение во временный файл
            temp_image_path = f'temp_image.{image_format}'
            with open(temp_image_path, 'wb') as img_file:
                img_file.write(image_data)
            # Добавляем изображение в новый документ
            target_doc.add_picture(temp_image_path, width=Inches(2))  # Задайте нужную ширину
            os.remove(temp_image_path)  # Удаляем временный файл

# Функция для копирования таблиц
def copy_tables(source_doc, target_doc):
    for table in source_doc.tables:
        new_table = target_doc.add_table(rows=0, cols=len(table.columns))
        for row in table.rows:
            new_row = new_table.add_row().cells
            for idx, cell in enumerate(row.cells):
                if cell is not None and cell.text is not None:
                    new_row[idx].text = cell.text.strip()

# Функция для копирования элементов (заголовков, параграфов, таблиц и изображений)
def copy_elements(source_doc, target_doc):
    for paragraph in source_doc.paragraphs:
        # Копируем заголовки
        if paragraph.style.name.startswith('Heading'):
            target_doc.add_heading(paragraph.text, level=int(paragraph.style.name[-1]))
        else:
            target_doc.add_paragraph(paragraph.text)

        # Проверяем следующий элемент, чтобы увидеть, является ли он таблицей
        next_element = paragraph._element.getnext()
        if next_element is not None:
            # Если следующий элемент - таблица
            if next_element.tag.endswith('tbl'):
                new_table = target_doc.add_table(rows=0, cols=len(next_element.xpath('.//w:tr[1]//w:tc')))
                for row in next_element.xpath('.//w:tr'):
                    new_row = new_table.add_row().cells
                    for idx, cell in enumerate(row.xpath('.//w:tc')):
                        if cell is not None and cell.text is not None:
                            new_row[idx].text = cell.text.strip()
                continue  # Пропускаем добавление новой строки

            # Если следующий элемент - изображение
            elif next_element.tag.endswith('blip'):
                # Получаем данные изображения
                image_data = next_element.blob
                image_format = next_element.content_type.split('/')[1]
                temp_image_path = f'temp_image.{image_format}'
                with open(temp_image_path, 'wb') as img_file:
                    img_file.write(image_data)
                target_doc.add_picture(temp_image_path, width=Inches(2))  # Задайте нужную ширину
                os.remove(temp_image_path)  # Удаляем временный файл

# Основная функция для переноса данных
def transfer_content(source_path, target_path):
    # Конвертация .doc в .docx
    converted_docx = convert_doc_to_docx(source_path)

    source_doc = Document(converted_docx)
    target_doc = Document()

    copy_elements(source_doc, target_doc)  # Копируем заголовки, текст и таблицы
    copy_images(source_doc, target_doc)  # Копируем изображения (если есть)

    target_doc.save(target_path)

# Указываем путь к исходному и целевому файлам
source_file_path = r"C:\Users\demchenko\Desktop\SRFauto\test\old_docs\1.doc"  # Ваш файл DOC
target_file_path = r"C:\Users\demchenko\Desktop\SRFauto\test\new_docs\output.docx"  # Имя нового DOCX файла

# Выполняем перенос
try:
    transfer_content(source_file_path, target_file_path)
    print("Перенос завершен успешно.")
except Exception as e:
    print(f"Произошла ошибка: {e}")
